import os
import telebot
from flask import Flask, request
from PIL import Image, ImageOps
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import threading

# –ü–æ–ª—É—á–∞–µ–º —Ç–æ–∫–µ–Ω –∏–∑ –ø–µ—Ä–µ–º–µ–Ω–Ω—ã—Ö –æ–∫—Ä—É–∂–µ–Ω–∏—è Railway
TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN', "8204855927:AAE6WxvaZl-kqM3zbSRql1J_dr1l1NteYeA")

bot = telebot.TeleBot(TOKEN)
app = Flask(__name__)
user_sessions = {}

# –í–∞—à —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π –∫–æ–¥ —Ñ—É–Ω–∫—Ü–∏–π (–æ—Å—Ç–∞–≤–ª—è–µ–º –±–µ–∑ –∏–∑–º–µ–Ω–µ–Ω–∏–π)
@bot.message_handler(commands=['start'])
def start(message):
    user_id = message.from_user.id
    if user_id not in user_sessions:
        user_sessions[user_id] = {'photos': [], 'format': 'pdf'}

    current_format = user_sessions[user_id]['format']
    format_name = "PDF" if current_format == 'pdf' else "DOCX"

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    btn_pdf = telebot.types.KeyboardButton('üìÑ PDF')
    btn_docx = telebot.types.KeyboardButton('üìù DOCX')
    btn_create = telebot.types.KeyboardButton('/create')
    btn_status = telebot.types.KeyboardButton('/status')
    markup.add(btn_pdf, btn_docx, btn_create, btn_status)

    bot.send_message(
        message.chat.id,
        f"üì∏ –ü—Ä–∏–≤–µ—Ç! –Ø –±–æ—Ç –¥–ª—è —Å–æ–∑–¥–∞–Ω–∏—è PDF –∏–ª–∏ DOCX –∏–∑ —Ñ–æ—Ç–æ.\n\n"
        f"üéØ –¢–µ–∫—É—â–∏–π —Ñ–æ—Ä–º–∞—Ç: {format_name}\n"
        f"üì∑ –§–æ—Ç–æ: {len(user_sessions[user_id]['photos'])}\n\n"
        "–ò—Å–ø–æ–ª—å–∑—É–π –∫–Ω–æ–ø–∫–∏ –¥–ª—è —É–ø—Ä–∞–≤–ª–µ–Ω–∏—è:",
        reply_markup=markup
    )

@bot.message_handler(commands=['help'])
def help_cmd(message):
    help_text = """
üìñ **–ö–æ–º–∞–Ω–¥—ã –±–æ—Ç–∞:**

/start - –ø–æ–∫–∞–∑–∞—Ç—å –º–µ–Ω—é
/help - –ø–æ–∫–∞–∑–∞—Ç—å —Å–ø—Ä–∞–≤–∫—É
/create - —Å–æ–∑–¥–∞—Ç—å –¥–æ–∫—É–º–µ–Ω—Ç
/clear - –æ—á–∏—Å—Ç–∏—Ç—å –≤—Å–µ —Ñ–æ—Ç–æ
/status - –ø–æ–∫–∞–∑–∞—Ç—å —Å—Ç–∞—Ç—É—Å
/reset - –ø–æ–ª–Ω—ã–π —Å–±—Ä–æ—Å

üñºÔ∏è **–ö–∞–∫ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å:**
1. –í—ã–±–µ—Ä–∏ —Ñ–æ—Ä–º–∞—Ç (PDF –∏–ª–∏ DOCX)
2. –û—Ç–ø—Ä–∞–≤–ª—è–π —Ñ–æ—Ç–æ (–ø–æ –æ–¥–Ω–æ–º—É –∏–ª–∏ –≥—Ä—É–ø–ø–æ–π)
3. –ù–∞–∂–º–∏ /create –¥–ª—è —Å–æ–∑–¥–∞–Ω–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–∞
4. –ü–æ–ª—É—á–∏ –≥–æ—Ç–æ–≤—ã–π —Ñ–∞–π–ª!

üí° **–û—Å–æ–±–µ–Ω–Ω–æ—Å—Ç–∏ —Ñ–æ—Ä–º–∞—Ç–æ–≤:**
‚Ä¢ üìÑ PDF - –æ—Ç–ª–∏—á–Ω–æ —Å–æ—Ö—Ä–∞–Ω—è–µ—Ç –∫–∞—á–µ—Å—Ç–≤–æ, —É–Ω–∏–≤–µ—Ä—Å–∞–ª—å–Ω—ã–π
‚Ä¢ üìù DOCX - –º–æ–∂–Ω–æ —Ä–µ–¥–∞–∫—Ç–∏—Ä–æ–≤–∞—Ç—å, –¥–æ–±–∞–≤–ª—è—Ç—å —Ç–µ–∫—Å—Ç
"""
    bot.send_message(message.chat.id, help_text)

@bot.message_handler(commands=['format'])
def choose_format(message):
    user_id = message.from_user.id
    if user_id not in user_sessions:
        user_sessions[user_id] = {'photos': [], 'format': 'pdf'}

    current_format = user_sessions[user_id]['format']
    current_format_name = "PDF" if current_format == 'pdf' else "DOCX"

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    btn_pdf = telebot.types.KeyboardButton('üìÑ PDF')
    btn_docx = telebot.types.KeyboardButton('üìù DOCX')
    btn_back = telebot.types.KeyboardButton('–ù–∞–∑–∞–¥')
    markup.add(btn_pdf, btn_docx, btn_back)

    bot.send_message(
        message.chat.id,
        f"üéØ –¢–µ–∫—É—â–∏–π —Ñ–æ—Ä–º–∞—Ç: {current_format_name}\n\n"
        f"–í—ã–±–µ—Ä–∏ –Ω–æ–≤—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–æ–∫—É–º–µ–Ω—Ç–∞:",
        reply_markup=markup
    )

@bot.message_handler(func=lambda message: message.text in ['üìÑ PDF', 'üìù DOCX', '–ù–∞–∑–∞–¥'])
def handle_format_choice(message):
    user_id = message.from_user.id
    if user_id not in user_sessions:
        user_sessions[user_id] = {'photos': [], 'format': 'pdf'}

    if message.text == '–ù–∞–∑–∞–¥':
        show_main_menu(message)
        return

    if message.text == 'üìÑ PDF':
        user_sessions[user_id]['format'] = 'pdf'
        format_name = "PDF"
    else:
        user_sessions[user_id]['format'] = 'docx'
        format_name = "DOCX"

    show_main_menu(message, f"‚úÖ –£—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω —Ñ–æ—Ä–º–∞—Ç: {format_name}")

def show_main_menu(message, additional_text=""):
    user_id = message.from_user.id
    if user_id not in user_sessions:
        user_sessions[user_id] = {'photos': [], 'format': 'pdf'}

    current_format = user_sessions[user_id]['format']
    format_name = "PDF" if current_format == 'pdf' else "DOCX"

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    btn_pdf = telebot.types.KeyboardButton('üìÑ PDF')
    btn_docx = telebot.types.KeyboardButton('üìù DOCX')
    btn_create = telebot.types.KeyboardButton('/create')
    btn_status = telebot.types.KeyboardButton('/status')
    markup.add(btn_pdf, btn_docx, btn_create, btn_status)

    text = f"üì∏ –ë–æ—Ç –¥–ª—è —Å–æ–∑–¥–∞–Ω–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –∏–∑ —Ñ–æ—Ç–æ\n\n"
    if additional_text:
        text += f"{additional_text}\n\n"
    text += f"üéØ –¢–µ–∫—É—â–∏–π —Ñ–æ—Ä–º–∞—Ç: {format_name}\n"
    text += f"üì∑ –§–æ—Ç–æ: {len(user_sessions[user_id]['photos'])}\n\n"
    text += "–ò—Å–ø–æ–ª—å–∑—É–π –∫–Ω–æ–ø–∫–∏ –¥–ª—è —É–ø—Ä–∞–≤–ª–µ–Ω–∏—è:"

    bot.send_message(message.chat.id, text, reply_markup=markup)

@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    user_id = message.from_user.id

    if user_id not in user_sessions:
        user_sessions[user_id] = {'photos': [], 'format': 'pdf'}

    file_info = bot.get_file(message.photo[-1].file_id)
    downloaded_file = bot.download_file(file_info.file_path)

    user_sessions[user_id]['photos'].append(downloaded_file)
    count = len(user_sessions[user_id]['photos'])
    format_name = "PDF" if user_sessions[user_id]['format'] == 'pdf' else "DOCX"

    bot.reply_to(
        message,
        f"‚úÖ –§–æ—Ç–æ {count} –ø–æ–ª—É—á–µ–Ω–æ!\n"
        f"–§–æ—Ä–º–∞—Ç: {format_name}\n\n"
        f"–û—Ç–ø—Ä–∞–≤—å—Ç–µ –µ—â—ë —Ñ–æ—Ç–æ –∏–ª–∏ /create –¥–ª—è —Å–æ–∑–¥–∞–Ω–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–∞"
    )

@bot.message_handler(commands=['create'])
def create_document(message):
    user_id = message.from_user.id

    if user_id not in user_sessions or not user_sessions[user_id]['photos']:
        bot.reply_to(message, "‚ùå –°–Ω–∞—á–∞–ª–∞ –æ—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ!")
        return

    try:
        bot.send_message(message.chat.id, "üîÑ –°–æ–∑–¥–∞—é –¥–æ–∫—É–º–µ–Ω—Ç...")

        format_type = user_sessions[user_id]['format']
        photos = user_sessions[user_id]['photos']

        if format_type == 'pdf':
            file_buffer = create_pdf(photos)
            file_name = "photos.pdf"
            caption = f"üìÑ –í–∞—à PDF —Ñ–∞–π–ª –≥–æ—Ç–æ–≤!\n–°—Ç—Ä–∞–Ω–∏—Ü: {len(photos)}"
        else:
            file_buffer = create_docx(photos)
            file_name = "photos.docx"
            caption = f"üìù –í–∞—à DOCX —Ñ–∞–π–ª –≥–æ—Ç–æ–≤!\n–°—Ç—Ä–∞–Ω–∏—Ü: {len(photos)}"

        bot.send_document(
            message.chat.id,
            file_buffer,
            visible_file_name=file_name,
            caption=caption
        )

        user_sessions[user_id]['photos'] = []

    except Exception as e:
        bot.reply_to(message, f"‚ùå –û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ–∑–¥–∞–Ω–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞: {e}")

def create_pdf(photos_bytes):
    """–°–æ–∑–¥–∞–µ—Ç PDF –∏–∑ —Å–ø–∏—Å–∫–∞ –±–∞–π—Ç–æ–≤ —Ñ–æ—Ç–æ"""
    images = []
    for photo_bytes in photos_bytes:
        image = Image.open(io.BytesIO(photo_bytes))

        try:
            image = ImageOps.exif_transpose(image)
        except:
            pass

        if image.mode != 'RGB':
            image = image.convert('RGB')
        images.append(image)

    pdf_buffer = io.BytesIO()

    if len(images) == 1:
        images[0].save(pdf_buffer, format='PDF', quality=95)
    else:
        images[0].save(
            pdf_buffer,
            format='PDF',
            save_all=True,
            append_images=images[1:],
            quality=95
        )

    pdf_buffer.seek(0)
    return pdf_buffer

def create_docx(photos_bytes):
    """–°–æ–∑–¥–∞–µ—Ç DOCX –¥–æ–∫—É–º–µ–Ω—Ç –∏–∑ —Å–ø–∏—Å–∫–∞ –±–∞–π—Ç–æ–≤ —Ñ–æ—Ç–æ —Å –∑–∞–ø–æ–ª–Ω–µ–Ω–∏–µ–º –≤—Å–µ–π —Å—Ç—Ä–∞–Ω–∏—Ü—ã"""
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    content_width = 7.5
    content_height = 10.0

    for i, photo_bytes in enumerate(photos_bytes):
        image_stream = io.BytesIO(photo_bytes)

        with Image.open(image_stream) as img:
            try:
                img = ImageOps.exif_transpose(img)
            except:
                pass

            img_width, img_height = img.size
            aspect_ratio = img_height / img_width
            page_aspect_ratio = content_height / content_width

            if aspect_ratio > page_aspect_ratio:
                calculated_height = Inches(content_height)
                calculated_width = Inches(content_height / aspect_ratio)
            else:
                calculated_width = Inches(content_width)
                calculated_height = Inches(content_width * aspect_ratio)

        image_stream.seek(0)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        run.add_picture(image_stream, width=calculated_width, height=calculated_height)

        if i < len(photos_bytes) - 1:
            doc.add_page_break()

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer

@bot.message_handler(commands=['clear'])
def clear_photos(message):
    user_id = message.from_user.id
    if user_id in user_sessions and user_sessions[user_id]['photos']:
        count = len(user_sessions[user_id]['photos'])
        user_sessions[user_id]['photos'] = []
        bot.reply_to(message, f"üóëÔ∏è –£–¥–∞–ª–µ–Ω–æ {count} —Ñ–æ—Ç–æ")
        show_main_menu(message)
    else:
        bot.reply_to(message, "‚ÑπÔ∏è –ù–µ—Ç —Ñ–æ—Ç–æ –¥–ª—è –æ—á–∏—Å—Ç–∫–∏")

@bot.message_handler(commands=['reset'])
def reset_session(message):
    user_id = message.from_user.id
    if user_id in user_sessions:
        count = len(user_sessions[user_id]['photos'])
        user_sessions[user_id] = {'photos': [], 'format': 'pdf'}
        bot.reply_to(message, f"üîÑ –°–µ—Å—Å–∏—è —Å–±—Ä–æ—à–µ–Ω–∞! –£–¥–∞–ª–µ–Ω–æ {count} —Ñ–æ—Ç–æ")
        show_main_menu(message)
    else:
        user_sessions[user_id] = {'photos': [], 'format': 'pdf'}
        bot.reply_to(message, "üîÑ –°–µ—Å—Å–∏—è —Å–æ–∑–¥–∞–Ω–∞!")

@bot.message_handler(commands=['status'])
def show_status(message):
    user_id = message.from_user.id
    if user_id in user_sessions:
        photos_count = len(user_sessions[user_id]['photos'])
        format_type = user_sessions[user_id]['format']
        format_name = "PDF" if format_type == 'pdf' else "DOCX"

        status_text = (
            f"üìä –°—Ç–∞—Ç—É—Å:\n"
            f"‚Ä¢ –§–æ—Ç–æ: {photos_count}\n"
            f"‚Ä¢ –§–æ—Ä–º–∞—Ç: {format_name}\n"
        )

        if photos_count > 0:
            status_text += f"\n‚úÖ –ì–æ—Ç–æ–≤ –∫ —Å–æ–∑–¥–∞–Ω–∏—é! –ò—Å–ø–æ–ª—å–∑—É–π /create"
        else:
            status_text += f"\nüì∏ –û—Ç–ø—Ä–∞–≤—å —Ñ–æ—Ç–æ —á—Ç–æ–±—ã –Ω–∞—á–∞—Ç—å"

        bot.reply_to(message, status_text)
    else:
        bot.reply_to(message, "‚ÑπÔ∏è –ù–∞—á–Ω–∏ —Å /start")

@bot.message_handler(func=lambda message: True)
def handle_other_messages(message):
    if message.text.startswith('/'):
        bot.reply_to(message, "‚ùå –ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –∫–æ–º–∞–Ω–¥–∞. –ò—Å–ø–æ–ª—å–∑—É–π /help –¥–ª—è —Å–ø—Ä–∞–≤–∫–∏")
    else:
        show_main_menu(message)

# –ù–æ–≤—ã–π –∫–æ–¥ –¥–ª—è Railway
@app.route('/')
def home():
    return "ü§ñ Telegram Bot is running! Use /start in Telegram."

@app.route('/health')
def health():
    return "OK"

@app.route('/webhook', methods=['POST'])
def webhook():
    if request.headers.get('content-type') == 'application/json':
        json_string = request.get_data().decode('utf-8')
        update = telebot.types.Update.de_json(json_string)
        bot.process_new_updates([update])
        return ''
    else:
        return 'Invalid content type', 403

def set_webhook():
    # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –≤–µ–±—Ö—É–∫ –¥–ª—è Railway
    webhook_url = f"https://{os.environ.get('RAILWAY_STATIC_URL', '')}/webhook"
    if webhook_url.startswith('https://'):
        bot.remove_webhook()
        bot.set_webhook(url=webhook_url)
        print(f"Webhook set to: {webhook_url}")
    else:
        print("Using polling mode")

def run_bot():
    """–ó–∞–ø—É—Å–∫–∞–µ—Ç –±–æ—Ç–∞ –≤ —Ä–µ–∂–∏–º–µ polling (–∫–∞–∫ –∑–∞–ø–∞—Å–Ω–æ–π –≤–∞—Ä–∏–∞–Ω—Ç)"""
    print("üöÄ –ë–æ—Ç –∑–∞–ø—É—â–µ–Ω –≤ —Ä–µ–∂–∏–º–µ polling!")
    print("üì∏ –§–æ—Ä–º–∞—Ç—ã: PDF –∏ DOCX")
    try:
        bot.infinity_polling()
    except Exception as e:
        print(f"‚ùå –û—à–∏–±–∫–∞ –≤ –±–æ—Ç–µ: {e}")

if __name__ == '__main__':
    # –ü—ã—Ç–∞–µ–º—Å—è —É—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –≤–µ–±—Ö—É–∫, –µ—Å–ª–∏ –¥–æ—Å—Ç—É–ø–µ–Ω URL
    set_webhook()
    
    # –ó–∞–ø—É—Å–∫–∞–µ–º Flask –ø—Ä–∏–ª–æ–∂–µ–Ω–∏–µ
    port = int(os.environ.get('PORT', 5000))
    print(f"Starting server on port {port}")
    app.run(host='0.0.0.0', port=port)
