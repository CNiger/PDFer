import os
import telebot
from flask import Flask, request
import threading
from PIL import Image, ImageOps
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time

# ÐŸÐ¾Ð»ÑƒÑ‡Ð°ÐµÐ¼ Ñ‚Ð¾ÐºÐµÐ½ Ð¸Ð· Ð¿ÐµÑ€ÐµÐ¼ÐµÐ½Ð½Ñ‹Ñ… Ð¾ÐºÑ€ÑƒÐ¶ÐµÐ½Ð¸Ñ Railway
TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN', "8204855927:AAE6WxvaZl-kqM3zbSRql1J_dr1l1NteYeA")

bot = telebot.TeleBot(TOKEN)
app = Flask(__name__)
user_sessions = {}

# Ð”Ð¾Ð±Ð°Ð²Ð»ÑÐµÐ¼ ÑÑ‡ÐµÑ‚Ñ‡Ð¸Ðº Ð´Ð»Ñ Ð¿Ð¾Ñ€ÑÐ´ÐºÐ° Ñ„Ð¾Ñ‚Ð¾
@bot.message_handler(commands=['start'])
def start(message):
    user_id = message.from_user.id
    if user_id not in user_sessions:
        user_sessions[user_id] = {'photos': [], 'format': 'pdf', 'photo_counter': 0}

    current_format = user_sessions[user_id]['format']
    format_name = "PDF" if current_format == 'pdf' else "DOCX"

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    btn_pdf = telebot.types.KeyboardButton('ðŸ“„ PDF')
    btn_docx = telebot.types.KeyboardButton('ðŸ“ DOCX')
    btn_create = telebot.types.KeyboardButton('/create')
    btn_status = telebot.types.KeyboardButton('/status')
    markup.add(btn_pdf, btn_docx, btn_create, btn_status)

    bot.send_message(
        message.chat.id,
        f"ðŸ“¸ ÐŸÑ€Ð¸Ð²ÐµÑ‚! Ð¯ Ð±Ð¾Ñ‚ Ð´Ð»Ñ ÑÐ¾Ð·Ð´Ð°Ð½Ð¸Ñ PDF Ð¸Ð»Ð¸ DOCX Ð¸Ð· Ñ„Ð¾Ñ‚Ð¾.\n\n"
        f"ðŸŽ¯ Ð¢ÐµÐºÑƒÑ‰Ð¸Ð¹ Ñ„Ð¾Ñ€Ð¼Ð°Ñ‚: {format_name}\n"
        f"ðŸ“· Ð¤Ð¾Ñ‚Ð¾: {len(user_sessions[user_id]['photos'])}\n\n"
        "Ð˜ÑÐ¿Ð¾Ð»ÑŒÐ·ÑƒÐ¹ ÐºÐ½Ð¾Ð¿ÐºÐ¸ Ð´Ð»Ñ ÑƒÐ¿Ñ€Ð°Ð²Ð»ÐµÐ½Ð¸Ñ:",
        reply_markup=markup
    )

@bot.message_handler(commands=['help'])
def help_cmd(message):
    help_text = """
ðŸ“– **ÐšÐ¾Ð¼Ð°Ð½Ð´Ñ‹ Ð±Ð¾Ñ‚Ð°:**

/start - Ð¿Ð¾ÐºÐ°Ð·Ð°Ñ‚ÑŒ Ð¼ÐµÐ½ÑŽ
/help - Ð¿Ð¾ÐºÐ°Ð·Ð°Ñ‚ÑŒ ÑÐ¿Ñ€Ð°Ð²ÐºÑƒ
/create - ÑÐ¾Ð·Ð´Ð°Ñ‚ÑŒ Ð´Ð¾ÐºÑƒÐ¼ÐµÐ½Ñ‚
/clear - Ð¾Ñ‡Ð¸ÑÑ‚Ð¸Ñ‚ÑŒ Ð²ÑÐµ Ñ„Ð¾Ñ‚Ð¾
/status - Ð¿Ð¾ÐºÐ°Ð·Ð°Ñ‚ÑŒ ÑÑ‚Ð°Ñ‚ÑƒÑ
/reset - Ð¿Ð¾Ð»Ð½Ñ‹Ð¹ ÑÐ±Ñ€Ð¾Ñ
"""
    bot.send_message(message.chat.id, help_text)

# Ð˜Ð·Ð¼ÐµÐ½ÑÐµÐ¼ Ð¾Ð±Ñ€Ð°Ð±Ð¾Ñ‚Ñ‡Ð¸Ðº Ñ„Ð¾Ñ‚Ð¾ - Ð´Ð¾Ð±Ð°Ð²Ð»ÑÐµÐ¼ Ð¿Ð¾Ñ€ÑÐ´ÐºÐ¾Ð²Ñ‹Ð¹ Ð½Ð¾Ð¼ÐµÑ€
@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    user_id = message.from_user.id

    if user_id not in user_sessions:
        user_sessions[user_id] = {'photos': [], 'format': 'pdf', 'photo_counter': 0}

    # Ð£Ð²ÐµÐ»Ð¸Ñ‡Ð¸Ð²Ð°ÐµÐ¼ ÑÑ‡ÐµÑ‚Ñ‡Ð¸Ðº Ð´Ð»Ñ ÐºÐ°Ð¶Ð´Ð¾Ð³Ð¾ Ð½Ð¾Ð²Ð¾Ð³Ð¾ Ñ„Ð¾Ñ‚Ð¾
    user_sessions[user_id]['photo_counter'] += 1
    order_number = user_sessions[user_id]['photo_counter']
    
    file_info = bot.get_file(message.photo[-1].file_id)
    downloaded_file = bot.download_file(file_info.file_path)

    # Ð¡Ð¾Ñ…Ñ€Ð°Ð½ÑÐµÐ¼ Ñ„Ð¾Ñ‚Ð¾ Ñ Ð¿Ð¾Ñ€ÑÐ´ÐºÐ¾Ð²Ñ‹Ð¼ Ð½Ð¾Ð¼ÐµÑ€Ð¾Ð¼ Ð¸ Ð²Ñ€ÐµÐ¼ÐµÐ½ÐµÐ¼ Ð¿Ð¾Ð»ÑƒÑ‡ÐµÐ½Ð¸Ñ
    user_sessions[user_id]['photos'].append({
        'data': downloaded_file,
        'order': order_number,
        'timestamp': time.time()
    })
    
    count = len(user_sessions[user_id]['photos'])
    format_name = "PDF" if user_sessions[user_id]['format'] == 'pdf' else "DOCX"

    bot.reply_to(
        message,
        f"âœ… Ð¤Ð¾Ñ‚Ð¾ {count} Ð¿Ð¾Ð»ÑƒÑ‡ÐµÐ½Ð¾!\n"
        f"Ð¤Ð¾Ñ€Ð¼Ð°Ñ‚: {format_name}\n\n"
        f"ÐžÑ‚Ð¿Ñ€Ð°Ð²ÑŒÑ‚Ðµ ÐµÑ‰Ñ‘ Ñ„Ð¾Ñ‚Ð¾ Ð¸Ð»Ð¸ /create Ð´Ð»Ñ ÑÐ¾Ð·Ð´Ð°Ð½Ð¸Ñ Ð´Ð¾ÐºÑƒÐ¼ÐµÐ½Ñ‚Ð°"
    )

@bot.message_handler(commands=['create'])
def create_document(message):
    user_id = message.from_user.id

    if user_id not in user_sessions or not user_sessions[user_id]['photos']:
        bot.reply_to(message, "âŒ Ð¡Ð½Ð°Ñ‡Ð°Ð»Ð° Ð¾Ñ‚Ð¿Ñ€Ð°Ð²ÑŒÑ‚Ðµ Ñ„Ð¾Ñ‚Ð¾!")
        return

    try:
        bot.send_message(message.chat.id, "ðŸ”„ Ð¡Ð¾Ð·Ð´Ð°ÑŽ Ð´Ð¾ÐºÑƒÐ¼ÐµÐ½Ñ‚...")

        format_type = user_sessions[user_id]['format']
        
        # Ð¡Ð¾Ñ€Ñ‚Ð¸Ñ€ÑƒÐµÐ¼ Ñ„Ð¾Ñ‚Ð¾ Ð¿Ð¾ Ð¿Ð¾Ñ€ÑÐ´ÐºÑƒ Ð´Ð¾Ð±Ð°Ð²Ð»ÐµÐ½Ð¸Ñ
        sorted_photos = sorted(user_sessions[user_id]['photos'], 
                             key=lambda x: x['order'])
        
        # Ð˜Ð·Ð²Ð»ÐµÐºÐ°ÐµÐ¼ Ñ‚Ð¾Ð»ÑŒÐºÐ¾ Ð´Ð°Ð½Ð½Ñ‹Ðµ Ñ„Ð¾Ñ‚Ð¾
        photos_data = [photo['data'] for photo in sorted_photos]

        if format_type == 'pdf':
            file_buffer = create_pdf(photos_data)
            file_name = "photos.pdf"
            caption = f"ðŸ“„ Ð’Ð°Ñˆ PDF Ñ„Ð°Ð¹Ð» Ð³Ð¾Ñ‚Ð¾Ð²!\nÐ¡Ñ‚Ñ€Ð°Ð½Ð¸Ñ†: {len(photos_data)}"
        else:
            file_buffer = create_docx(photos_data)
            file_name = "photos.docx"
            caption = f"ðŸ“ Ð’Ð°Ñˆ DOCX Ñ„Ð°Ð¹Ð» Ð³Ð¾Ñ‚Ð¾Ð²!\nÐ¡Ñ‚Ñ€Ð°Ð½Ð¸Ñ†: {len(photos_data)}"

        bot.send_document(
            message.chat.id,
            file_buffer,
            visible_file_name=file_name,
            caption=caption
        )

        user_sessions[user_id]['photos'] = []
        user_sessions[user_id]['photo_counter'] = 0  # Ð¡Ð±Ñ€Ð°ÑÑ‹Ð²Ð°ÐµÐ¼ ÑÑ‡ÐµÑ‚Ñ‡Ð¸Ðº

    except Exception as e:
        bot.reply_to(message, f"âŒ ÐžÑˆÐ¸Ð±ÐºÐ° Ð¿Ñ€Ð¸ ÑÐ¾Ð·Ð´Ð°Ð½Ð¸Ð¸ Ð´Ð¾ÐºÑƒÐ¼ÐµÐ½Ñ‚Ð°: {e}")

# ÐžÑÑ‚Ð°Ð»ÑŒÐ½Ñ‹Ðµ Ñ„ÑƒÐ½ÐºÑ†Ð¸Ð¸ Ð¾ÑÑ‚Ð°ÑŽÑ‚ÑÑ Ð±ÐµÐ· Ð¸Ð·Ð¼ÐµÐ½ÐµÐ½Ð¸Ð¹
def create_pdf(photos_bytes):
    images = []
    for photo_bytes in photos_bytes:
        image = Image.open(io.BytesIO(photo_bytes))
        try:
            image = ImageOps.exif_transpose(image)
        except:
            pass
        if image.mode != 'RGB':
            image = image.convert('RGB')
        images.append(image)

    pdf_buffer = io.BytesIO()
    if len(images) == 1:
        images[0].save(pdf_buffer, format='PDF', quality=95)
    else:
        images[0].save(
            pdf_buffer,
            format='PDF',
            save_all=True,
            append_images=images[1:],
            quality=95
        )
    pdf_buffer.seek(0)
    return pdf_buffer

def create_docx(photos_bytes):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    content_width = 7.5
    content_height = 10.0

    for i, photo_bytes in enumerate(photos_bytes):
        image_stream = io.BytesIO(photo_bytes)
        with Image.open(image_stream) as img:
            try:
                img = ImageOps.exif_transpose(img)
            except:
                pass
            img_width, img_height = img.size
            aspect_ratio = img_height / img_width
            page_aspect_ratio = content_height / content_width

            if aspect_ratio > page_aspect_ratio:
                calculated_height = Inches(content_height)
                calculated_width = Inches(content_height / aspect_ratio)
            else:
                calculated_width = Inches(content_width)
                calculated_height = Inches(content_width * aspect_ratio)

        image_stream.seek(0)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_stream, width=calculated_width, height=calculated_height)

        if i < len(photos_bytes) - 1:
            doc.add_page_break()

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

@bot.message_handler(commands=['clear'])
def clear_photos(message):
    user_id = message.from_user.id
    if user_id in user_sessions and user_sessions[user_id]['photos']:
        count = len(user_sessions[user_id]['photos'])
        user_sessions[user_id]['photos'] = []
        user_sessions[user_id]['photo_counter'] = 0  # Ð¡Ð±Ñ€Ð°ÑÑ‹Ð²Ð°ÐµÐ¼ ÑÑ‡ÐµÑ‚Ñ‡Ð¸Ðº
        bot.reply_to(message, f"ðŸ—‘ï¸ Ð£Ð´Ð°Ð»ÐµÐ½Ð¾ {count} Ñ„Ð¾Ñ‚Ð¾")
    else:
        bot.reply_to(message, "â„¹ï¸ ÐÐµÑ‚ Ñ„Ð¾Ñ‚Ð¾ Ð´Ð»Ñ Ð¾Ñ‡Ð¸ÑÑ‚ÐºÐ¸")

@bot.message_handler(commands=['status'])
def show_status(message):
    user_id = message.from_user.id
    if user_id in user_sessions:
        photos_count = len(user_sessions[user_id]['photos'])
        format_type = user_sessions[user_id]['format']
        format_name = "PDF" if format_type == 'pdf' else "DOCX"
        status_text = f"ðŸ“Š Ð¡Ñ‚Ð°Ñ‚ÑƒÑ:\nâ€¢ Ð¤Ð¾Ñ‚Ð¾: {photos_count}\nâ€¢ Ð¤Ð¾Ñ€Ð¼Ð°Ñ‚: {format_name}"
        bot.reply_to(message, status_text)
    else:
        bot.reply_to(message, "â„¹ï¸ ÐÐ°Ñ‡Ð½Ð¸ Ñ /start")

# Ð”Ð¾Ð±Ð°Ð²Ð»ÑÐµÐ¼ Ð¾Ð±Ñ€Ð°Ð±Ð¾Ñ‚Ñ‡Ð¸Ðº Ð´Ð»Ñ ÐºÐ½Ð¾Ð¿Ð¾Ðº Ñ„Ð¾Ñ€Ð¼Ð°Ñ‚Ð°
@bot.message_handler(func=lambda message: message.text in ['ðŸ“„ PDF', 'ðŸ“ DOCX'])
def handle_format_buttons(message):
    user_id = message.from_user.id
    if user_id not in user_sessions:
        user_sessions[user_id] = {'photos': [], 'format': 'pdf', 'photo_counter': 0}
    
    if message.text == 'ðŸ“„ PDF':
        user_sessions[user_id]['format'] = 'pdf'
        format_name = "PDF"
    else:
        user_sessions[user_id]['format'] = 'docx'
        format_name = "DOCX"
    
    bot.reply_to(message, f"âœ… Ð¤Ð¾Ñ€Ð¼Ð°Ñ‚ Ð¸Ð·Ð¼ÐµÐ½ÐµÐ½ Ð½Ð° {format_name}")

# Flask Ð¼Ð°Ñ€ÑˆÑ€ÑƒÑ‚Ñ‹ Ð´Ð»Ñ Railway (Ð±ÐµÐ· Ð¸Ð·Ð¼ÐµÐ½ÐµÐ½Ð¸Ð¹)
@app.route('/')
def home():
    return "ðŸ¤– Telegram Bot is running! Use /start in Telegram."

@app.route('/webhook', methods=['POST'])
def webhook():
    if request.headers.get('content-type') == 'application/json':
        json_string = request.get_data().decode('utf-8')
        update = telebot.types.Update.de_json(json_string)
        bot.process_new_updates([update])
        return ''
    else:
        return 'Invalid content type', 403

# Ð¤ÑƒÐ½ÐºÑ†Ð¸Ñ Ð´Ð»Ñ ÑƒÑÑ‚Ð°Ð½Ð¾Ð²ÐºÐ¸ Ð²ÐµÐ±Ñ…ÑƒÐºÐ° (Ð±ÐµÐ· Ð¸Ð·Ð¼ÐµÐ½ÐµÐ½Ð¸Ð¹)
def set_webhook():
    try:
        railway_url = os.environ.get('RAILWAY_STATIC_URL')
        if railway_url:
            webhook_url = f"{railway_url}/webhook"
            bot.remove_webhook()
            bot.set_webhook(url=webhook_url)
            print(f"âœ… Webhook ÑƒÑÑ‚Ð°Ð½Ð¾Ð²Ð»ÐµÐ½: {webhook_url}")
        else:
            print("â„¹ï¸ RAILWAY_STATIC_URL Ð½Ðµ Ð½Ð°Ð¹Ð´ÐµÐ½, Ð¸ÑÐ¿Ð¾Ð»ÑŒÐ·ÑƒÐµÐ¼ polling")
            threading.Thread(target=run_polling, daemon=True).start()
    except Exception as e:
        print(f"âŒ ÐžÑˆÐ¸Ð±ÐºÐ° ÑƒÑÑ‚Ð°Ð½Ð¾Ð²ÐºÐ¸ webhook: {e}")
        threading.Thread(target=run_polling, daemon=True).start()

def run_polling():
    print("ðŸ”„ Ð—Ð°Ð¿ÑƒÑÐºÐ°ÐµÐ¼ Ð±Ð¾Ñ‚Ð° Ð² Ñ€ÐµÐ¶Ð¸Ð¼Ðµ polling...")
    try:
        bot.infinity_polling()
    except Exception as e:
        print(f"âŒ ÐžÑˆÐ¸Ð±ÐºÐ° Ð² polling: {e}")

# ÐŸÑ€Ð¸ Ð·Ð°Ð¿ÑƒÑÐºÐµ Ð¿Ñ€Ð¸Ð»Ð¾Ð¶ÐµÐ½Ð¸Ñ (Ð±ÐµÐ· Ð¸Ð·Ð¼ÐµÐ½ÐµÐ½Ð¸Ð¹)
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    print(f"ðŸš€ Starting server on port {port}")
    set_webhook()
    app.run(host='0.0.0.0', port=port)
else:
    set_webhook()
